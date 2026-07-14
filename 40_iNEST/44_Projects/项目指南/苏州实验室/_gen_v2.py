# -*- coding: utf-8 -*-
import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(BASE, 'temp_guide_v4.md')
OUT = os.path.join(BASE, 'iNEST_Project_Guide_Final.docx')

with open(MD, 'r', encoding='utf-8') as f:
    content = f.read()

content = re.sub(r'### 1\\.3 .*?(?=###|\\Z)', '', content, flags=re.DOTALL)
content = re.sub(r'.*?V9.*?V32.*?\\n', '', content)
content = re.sub(r'.*?SDI.*?experiment.*?\\n', '', content)
content = re.sub(r'\\n{4,}', '\\n\\n\\n', content)

lines = content.split('\\n')
doc = Document()

for sec in doc.sections:
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

sty = doc.styles['Normal']
sty.font.name = 'Calibri'
sty.font.size = Pt(11)
sty.paragraph_format.space_after = Pt(6)
sty.paragraph_format.line_spacing = 1.10

for lvl, sz, clr in [(1,16,'2E74B5'),(2,13,'2E74B5'),(3,12,'1F4D78')]:
    hs = doc.styles['Heading %d' % lvl]
    hs.font.name = 'Calibri'
    hs.font.size = Pt(sz)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor.from_string(clr)
    hs.paragraph_format.space_before = Pt({1:16,2:12,3:8}[lvl])
    hs.paragraph_format.space_after = Pt({1:8,2:6,3:4}[lvl])

def ar(p, t, bold=False, italic=False, sz=None, clr=None):
    if not t: return
    r = p.add_run(t)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Calibri'
    if sz: r.font.size = Pt(sz)
    if clr: r.font.color.rgb = RGBColor.from_string(clr)

def ah(t, lv):
    h = doc.add_heading(t, level=lv)
    for r in h.runs: r.font.name = 'Calibri'

def pf(text):
    parts = []
    rest = text
    while rest:
        m = re.search(r'\\*\\*(.+?)\\*\\*', rest)
        if not m:
            parts.append(('n', rest))
            break
        if m.start() > 0:
            parts.append(('n', rest[:m.start()]))
        parts.append(('b', m.group(1)))
        rest = rest[m.end():]
    return parts

def ap(text, quote=False, bullet=False):
    if bullet:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph()
    if quote:
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.right_indent = Cm(1.0)
    for k, t in pf(text):
        if quote:
            ar(p, t, bold=(k=='b'), italic=True, sz=10, clr='2E74B5')
        else:
            ar(p, t, bold=(k=='b'))

def mt(headers, rows):
    n = len(headers)
    ns = nsdecls('w')
    t = doc.add_table(rows=len(rows)+1, cols=n, style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h_text in enumerate(headers):
        c = t.cell(0, ci)
        c.text = ''
        for k, txt in pf(h_text):
            ar(c.paragraphs[0], txt, bold=True, sz=9, clr='FFFFFF')
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        xml = '<w:shd %s w:fill="2E74B5" w:val="clear"/>' % ns
        c._element.get_or_add_tcPr().append(parse_xml(xml))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci < n:
                c = t.cell(ri+1, ci)
                c.text = ''
                for k, txt in pf(str(val)):
                    ar(c.paragraphs[0], txt, bold=(k=='b'), sz=8.5)
                if ri % 2 == 1:
                    xml2 = '<w:shd %s w:fill="F2F6FC" w:val="clear"/>' % ns
                    c._element.get_or_add_tcPr().append(parse_xml(xml2))
    doc.add_paragraph()

i = 0
in_table = False
tl = []
in_code = False

while i < len(lines):
    line = lines[i]
    s = line.strip()
    if i == 0 and s == '---':
        i += 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1
        continue
    if s.startswith('`' + '``'):
        in_code = not in_code
        i += 1
        continue
    if in_code:
        i += 1
        continue
    if not s:
        i += 1
        continue
    if s.startswith('# ') and not in_table:
        ah(s[2:], 1)
        i += 1
        continue
    if s.startswith('## '):
        ah(s[3:], 2)
        i += 1
        continue
    if s.startswith('### '):
        ah(s[4:], 3)
        i += 1
        continue
    if s.startswith('|') and not in_table:
        in_table = True
        tl = [s]
        i += 1
        continue
    if in_table:
        if s.startswith('|'):
            tl.append(s)
            i += 1
            continue
        else:
            rows = []
            for l in tl:
                if not l.startswith('|---'):
                    cells = [c.strip() for c in l.split('|')[1:-1]]
                    ok = True
                    for c in cells:
                        if re.match(r'^-+$', c):
                            ok = False
                            break
                    if ok and cells:
                        rows.append(cells)
            if rows and len(rows) > 1:
                mt(rows[0], rows[1:])
            tl = []
            in_table = False
            continue
    if s.startswith('>'):
        ap(s.lstrip('> ').strip(), quote=True)
        i += 1
        continue
    if s == '---':
        p = doc.add_paragraph(chr(0x2500) * 60)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue
    if s.startswith('- ') or s.startswith('* '):
        ap(s.lstrip('- *').strip(), bullet=True)
        i += 1
        continue
    if re.match(r'^\\d+\\.\\s', s):
        p = doc.add_paragraph(style='List Number')
        text = re.sub(r'^\\d+\\.\\s', '', s)
        for k, txt in pf(text):
            ar(p, txt, bold=(k=='b'))
        i += 1
        continue
    ap(s)
    i += 1

for sec in doc.sections:
    f = sec.footer
    f.is_linked_to_previous = False
    fp = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar(fp, 'Suzhou National Materials Lab  |  iNEST Pilot Project Guide  |  Confidential', sz=8, clr='888888')

doc.save(OUT)
print('OK: %.1f KB' % (os.path.getsize(OUT)/1024))